import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_document():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles Setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Title Page ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(100)
    title_run = title_p.add_run("DGEC PROJECT CONTROL DASHBOARD")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Deep Navy

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(200)
    sub_run = subtitle_p.add_run("System Architecture & User Workflow Specification")
    sub_run.font.size = Pt(16)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x5C, 0x76, 0x8D) # Muted Teal

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.add_run("Prepared for: ").bold = True
    info_p.add_run("DGEC Management\n")
    info_p.add_run("Version: ").bold = True
    info_p.add_run("1.0.0\n")
    info_p.add_run("Date: ").bold = True
    info_p.add_run("July 30, 2026\n")

    doc.add_page_break()

    # --- Table of Contents Placeholder ---
    h1 = doc.add_heading("Table of Contents", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h1.paragraph_format.space_after = Pt(12)

    toc_p = doc.add_paragraph()
    toc_p.add_run("1. Executive Summary ....................................................................................... 3\n")
    toc_p.add_run("2. System Architecture & Tech Stack ................................................................ 3\n")
    toc_p.add_run("3. User Roles & Workflow Specifications ............................................................. 4\n")
    toc_p.add_run("   3.1. Project Manager User Flow .................................................................... 4\n")
    toc_p.add_run("   3.2. Administrator User Flow ........................................................................ 5\n")
    toc_p.add_run("   3.3. Staff Engineer User Flow ....................................................................... 6\n")
    toc_p.add_run("   3.4. Client User Flow .................................................................................. 7\n")
    toc_p.add_run("4. Database Schema & Data Model ................................................................... 8\n")
    toc_p.add_run("5. Security & Deployment Best Practices ............................................................. 9\n")
    
    doc.add_page_break()

    # --- Section 1: Executive Summary ---
    h2 = doc.add_heading("1. Executive Summary", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    p1 = doc.add_paragraph(
        "The DGEC Project Control Dashboard is an enterprise-grade project management system "
        "designed specifically for engineering firms to streamline project oversight, task management, "
        "billing, staff assignment, and client transparency. The system ensures data integrity "
        "by utilizing secure authentication, role-based page redirects, and comprehensive activity "
        "logging."
    )
    p1.paragraph_format.space_after = Pt(12)

    # --- Section 2: System Architecture ---
    h3 = doc.add_heading("2. System Architecture & Tech Stack", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)

    doc.add_paragraph("The application is divided into a decoupled Frontend and Backend structure:")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("Frontend Web Application: ").bold = True
    bp1.add_run("Built using React and Vite, styled with custom responsive CSS. Built static assets are served from webapp.greensparrows.com.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("Backend API Server: ").bold = True
    bp2.add_run("Built using Node.js and Express.js. Handles business logic, query routing, mail delivery via NodeMailer, and PDF invoice generation via PDFKit. Hosted at api2.greensparrows.com.")

    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.add_run("Relational Database: ").bold = True
    bp3.add_run("MySQL database hosting tables with relational links to manage entities (projects, tasks, invoices, history, etc.).")

    # --- Section 3: User Roles & Workflows ---
    doc.add_page_break()
    h4 = doc.add_heading("3. User Roles & Workflow Specifications", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(12)

    # 3.1 Project Manager
    sh1 = doc.add_heading("3.1. Project Manager User Flow", level=2)
    sh1.runs[0].font.color.rgb = RGBColor(0x5C, 0x76, 0x8D)
    sh1.paragraph_format.space_before = Pt(12)
    
    p_pm = doc.add_paragraph(
        "The Project Manager (PM) is the core operational user of the system. "
        "They oversee the entire engineering lifecycle, including client relations, project initialization, "
        "document tracking, teammate assignments, and invoicing. They access the root URL (/)."
    )
    p_pm.paragraph_format.space_after = Pt(6)

    pm_flow = [
        "Login: Authenticates via login.html. Redirected to main control dashboard (index.html).",
        "Dashboard Overview: Views high-level stats (Total Projects, Open Tasks, Outstanding Invoices, Total Billing). Inspects real-time project progress bars.",
        "Project Management: Creates new projects, specifies sector, disciplines involved (e.g., Structure, HVAC, Electrical), and dates. Modifies status (Active, Concept, On Hold, Closed).",
        "Task Management: Creates sub-tasks under projects, sets completion percentages, targets dates, and assigns them to specific Staff Engineers.",
        "Team Assignments: Adds or removes team members from the active project workspace.",
        "Invoicing & Billing: Generates and records invoices, tracks payment status, and generates PDF receipts.",
        "Audit Log Review: Monitors the historical actions log to see which staff member modified what details and when."
    ]
    for step in pm_flow:
        sp = doc.add_paragraph(style='List Number')
        parts = step.split(": ")
        sp.add_run(parts[0] + ": ").bold = True
        sp.add_run(parts[1])

    # 3.2 Administrator
    sh2 = doc.add_heading("3.2. Administrator User Flow", level=2)
    sh2.runs[0].font.color.rgb = RGBColor(0x5C, 0x76, 0x8D)
    sh2.paragraph_format.space_before = Pt(12)

    p_admin = doc.add_paragraph(
        "The Administrator handles system settings, user database directory maintenance, and "
        "system-wide configuration. They access the administration panel (admin.html)."
    )
    p_admin.paragraph_format.space_after = Pt(6)

    admin_flow = [
        "Login: Authenticates via login.html using admin credentials. Redirected to admin.html.",
        "Directory Management: Adds, updates, or deletes users in the system (Staff, Client, PM accounts).",
        "Client Management: Registers new client companies, creates client usernames, and sets secure access passwords.",
        "Billing Auditing: Reviews, overrides, or deletes erroneous invoices from the database.",
        "System Settings: Edits dashboard dropdown configuration categories (Disciplines list, Project Statuses, Task categories)."
    ]
    for step in admin_flow:
        sp = doc.add_paragraph(style='List Number')
        parts = step.split(": ")
        sp.add_run(parts[0] + ": ").bold = True
        sp.add_run(parts[1])

    # 3.3 Staff Engineer
    sh3 = doc.add_heading("3.3. Staff Engineer User Flow", level=2)
    sh3.runs[0].font.color.rgb = RGBColor(0x5C, 0x76, 0x8D)
    sh3.paragraph_format.space_before = Pt(12)

    p_staff = doc.add_paragraph(
        "Staff Engineers are technical department heads (e.g. Senior Structural Engineer, HVAC Specialist) "
        "who execute project tasks. They access the staff portal (staff.html)."
    )
    p_staff.paragraph_format.space_after = Pt(6)

    staff_flow = [
        "Login: Authenticates via login.html. Redirected to staff.html.",
        "Task Queue: Inspects a filtered board showing only tasks assigned to them or their department.",
        "Progress Updates: Updates task progress percentages (0-100%) and sets statuses (In Progress, On Hold, Done).",
        "Document Review: Accesses document numbers associated with projects to coordinate blueprints and structural calculations.",
        "Activity Logging: Updates save changes, triggering an automatic record in the history audit log."
    ]
    for step in staff_flow:
        sp = doc.add_paragraph(style='List Number')
        parts = step.split(": ")
        sp.add_run(parts[0] + ": ").bold = True
        sp.add_run(parts[1])

    # 3.4 Client
    sh4 = doc.add_heading("3.4. Client User Flow", level=2)
    sh4.runs[0].font.color.rgb = RGBColor(0x5C, 0x76, 0x8D)
    sh4.paragraph_format.space_before = Pt(12)

    p_client = doc.add_paragraph(
        "Clients (e.g., developers, municipalities) are external stakeholders. They have read-only access "
        "to track progress and manage their financials. They access the client portal (client.html)."
    )
    p_client.paragraph_format.space_after = Pt(6)

    client_flow = [
        "Login: Authenticates via login.html (or client-login). Redirected to client.html.",
        "Project Tracking: Monitors overall project progress bars, active phases, and target completion dates.",
        "Invoices & Payments: Reviews sent invoices, tracks payment receipts, and downloads generated PDF statements.",
        "Document Registry: Views registered drawing and calculations numbers (AOR, structural blueprints) for submission tracking."
    ]
    for step in client_flow:
        sp = doc.add_paragraph(style='List Number')
        parts = step.split(": ")
        sp.add_run(parts[0] + ": ").bold = True
        sp.add_run(parts[1])

    # --- Section 4: Database Schema ---
    doc.add_page_break()
    h5 = doc.add_heading("4. Database Schema & Data Model", level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(6)

    p_db = doc.add_paragraph(
        "The system runs on a relational MySQL schema. Below are the key tables and their relationships:"
    )
    p_db.paragraph_format.space_after = Pt(12)

    # Database Table summary
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Table Name'
    hdr_cells[1].text = 'Primary Keys / Links'
    hdr_cells[2].text = 'Description'
    
    # Styles for table headers
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Background shading using XML manipulation for simple table style
        tcPr = cell._tc.get_or_add_tcPr()
        from docx.oxml import parse_xml
        shading = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="1B365D"/>')
        tcPr.append(shading)

    db_tables = [
        ("admin", "id (PK), uuid", "System administrators login accounts credentials."),
        ("clients", "id (PK), uuid", "Registered client companies and dashboard portal users."),
        ("projects", "id (PK), uuid, client_id (FK -> clients)", "Contains project schedules, sectors, statuses, and associated documents."),
        ("tasks", "id (PK), uuid, project_id (FK -> projects)", "Project workflow items assigned to Staff Engineers."),
        ("teammates", "id (PK), uuid, project_id (FK -> projects)", "Mapping of assigned engineers working on a project."),
        ("staff", "id (PK), uuid", "Employee directory containing names, emails, contact details, and department roles."),
        ("invoices", "id (PK), project_id (FK -> projects)", "Financial invoices, amounts, billing descriptions, and payment status."),
        ("history", "id (PK)", "Chronological audit logs of edit operations carried out across the panels.")
    ]

    for name, keys, desc in db_tables:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = keys
        row_cells[2].text = desc

    # Space after table
    doc.add_paragraph().paragraph_format.space_before = Pt(18)

    # --- Section 5: Security & Deployment ---
    h6 = doc.add_heading("5. Security & Deployment Best Practices", level=1)
    h6.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h6.paragraph_format.space_before = Pt(18)
    h6.paragraph_format.space_after = Pt(6)

    sec_p = doc.add_paragraph(
        "To safeguard credentials, the application implements the following deployment strategies:"
    )
    
    sec_points = [
        "In-Memory Environment Fallbacks: The backend defaults configuration options securely in-memory on Hostinger. This keeps actual passwords out of the public Git repository.",
        "Cross-Origin Resource Sharing (CORS): Configured in server.js to explicitly whitelist the webapp frontend origin, blocking malicious remote domains.",
        "Password Hashing: Authenticated accounts utilize bcrypt (with salt rounds) to hash and verify access passwords securely."
    ]
    for pt in sec_points:
        sp = doc.add_paragraph(style='List Bullet')
        parts = pt.split(": ")
        sp.add_run(parts[0] + ": ").bold = True
        sp.add_run(parts[1])

    # Save to workspace path
    workspace_path = "DGEC_Project_Document.docx"
    doc.save(workspace_path)
    
    # Save to artifacts path
    artifact_dir = "C:\\Users\\prave\\.gemini\\antigravity\\brain\\5c3aa846-18ce-4bf8-bbf7-ec1d57b520b8"
    if os.path.exists(artifact_dir):
        doc.save(os.path.join(artifact_dir, "DGEC_Project_Document.docx"))

    print("Success! Word document created.")

if __name__ == "__main__":
    create_document()
